import os
import subprocess
import shutil
from datetime import datetime

import pandas as pd
import tkinter as tk
from tkinter import ttk, messagebox, filedialog

# S'assure que le module openpyxl est disponible pour la gestion des fichiers Excel
try:
    import openpyxl  # noqa: F401
except ModuleNotFoundError:
    raise SystemExit(
        "Le module 'openpyxl' est requis. Installez-le avec 'pip install openpyxl'."
    )

# Modules utilises pour generer les fiches d'affectation
try:
    from docx import Document
    from docx2pdf import convert
except ModuleNotFoundError:
    raise SystemExit(
        "Les modules 'python-docx' et 'docx2pdf' sont requis. Installez-les avec 'pip install python-docx docx2pdf'."
    )

# Fichiers Excel utilises par l'application
TABLETTES_FILE = 'tablettes.xlsx'
AFFECT_FILE = 'affectations.xlsx'
INCIDENT_FILE = 'incidents.xlsx'

# Template Word et dossier de sortie pour les fiches
TEMPLATE_FILE = 'Fiche_Affectation_Materiel.docx'
FICHES_DIR = 'fiches'

# Valeurs possibles pour le statut d'une tablette
STATUS_OPTIONS = ['En stock', 'Affectée', 'En réparation', 'Perdue']


def init_files():
    """Cree les fichiers Excel avec les entetes s'ils n'existent pas."""
    if not os.path.exists(TABLETTES_FILE):
        df = pd.DataFrame(
            columns=[
                'N° Tablette',
                'Statut actuel',
                'Chargeur',
                'Powerbank',
                'Observations',
            ]
        )
        df.to_excel(TABLETTES_FILE, index=False, engine='openpyxl')
    if not os.path.exists(AFFECT_FILE):
        df = pd.DataFrame(
            columns=[
                "Date d'affectation",
                'N° Tablette',
                'Nom bénéficiaire',
                'Identifiant bénéficiaire',
            ]
        )
        df.to_excel(AFFECT_FILE, index=False, engine='openpyxl')
    if not os.path.exists(INCIDENT_FILE):
        df = pd.DataFrame(columns=['Date', 'N° Tablette', 'Nature incident', 'Déclarant', 'Lieu'])
        df.to_excel(INCIDENT_FILE, index=False, engine='openpyxl')


def load_tablettes():
    df = pd.read_excel(TABLETTES_FILE, engine='openpyxl')
    if 'Chargeur' not in df.columns:
        df['Chargeur'] = ''
    if 'Powerbank' not in df.columns:
        df['Powerbank'] = ''
    return df


def save_tablettes(df):
    df.to_excel(TABLETTES_FILE, index=False, engine='openpyxl')


def load_affectations():
    df = pd.read_excel(AFFECT_FILE, engine='openpyxl')
    if 'Identifiant bénéficiaire' not in df.columns:
        df['Identifiant bénéficiaire'] = ''
    return df


def save_affectations(df):
    df.to_excel(AFFECT_FILE, index=False, engine='openpyxl')


def load_incidents():
    return pd.read_excel(INCIDENT_FILE, engine='openpyxl')


def save_incidents(df):
    df.to_excel(INCIDENT_FILE, index=False, engine='openpyxl')


def _replace_placeholders(element, mapping):
    """Remplace les balises dans un element de document.

    Cette implementation remplace les placeholders au niveau du texte complet
    du paragraphe. Cela permet de gerer correctement les placeholders coupes en
    plusieurs runs par Word.
    """
    for paragraph in element.paragraphs:
        text = paragraph.text
        replaced = False
        for placeholder, value in mapping.items():
            if placeholder in text:
                text = text.replace(placeholder, value)
                replaced = True
        if replaced:
            # Efface les runs existants et insere le texte remplace dans un
            # nouveau run pour conserver un minimum de mise en forme.
            for run in paragraph.runs:
                run.text = ''
            paragraph.add_run(text)

    for table in element.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_placeholders(cell, mapping)


def generer_fiche(mapping, beneficiaire):
    """Cree une fiche d'affectation Word et son PDF."""
    if not os.path.exists(TEMPLATE_FILE):
        messagebox.showerror(
            'Modèle manquant',
            f"Le fichier {TEMPLATE_FILE} est introuvable.",
        )
        return
    os.makedirs(FICHES_DIR, exist_ok=True)
    doc = Document(TEMPLATE_FILE)
    placeholders = {f"{{{{{k}}}}}": str(v) for k, v in mapping.items()}
    placeholders['{{Date remise}}'] = ''
    _replace_placeholders(doc, placeholders)
    safe_name = beneficiaire.replace(' ', '_')
    doc_path = os.path.join(FICHES_DIR, f"fiche_{safe_name}.docx")
    pdf_path = os.path.join(FICHES_DIR, f"fiche_{safe_name}.pdf")
    doc.save(doc_path)
    try:
        convert(doc_path, pdf_path)
    except Exception as exc:  # pragma: no cover - dependance externe
        print(f'Erreur conversion PDF: {exc}')
    return doc_path, pdf_path


def ouvrir_dossier_fiches():
    """Ouvre le dossier contenant les fiches generées."""
    os.makedirs(FICHES_DIR, exist_ok=True)
    if os.name == 'nt':
        os.startfile(FICHES_DIR)
    else:
        subprocess.run(['xdg-open', FICHES_DIR])


def display_dataframe(df, title):
    """Affiche un DataFrame dans une nouvelle fen\u00eatre."""
    win = tk.Toplevel(root)
    win.title(title)
    tree = ttk.Treeview(win, columns=list(df.columns), show='headings')
    for col in df.columns:
        tree.heading(col, text=col)
        tree.column(col, width=120)
    for _, row in df.iterrows():
        tree.insert('', tk.END, values=list(row))
    tree.pack(expand=True, fill='both')


def open_tablettes():
    df = load_tablettes()
    display_dataframe(df, 'Tablettes')


def open_affectations():
    df = load_affectations()
    display_dataframe(df, 'Affectations')


def open_incidents():
    df = load_incidents()
    display_dataframe(df, 'Incidents')


def importer_tablettes():
    """Importe un fichier tablettes existant."""
    path = filedialog.askopenfilename(title='Sélectionnez le fichier tablettes',
                                      filetypes=[('Fichiers Excel', '*.xlsx')])
    if path:
        df = pd.read_excel(path, engine='openpyxl')

        # Vérifie la présence des colonnes requises
        required = ['N° Tablette', 'Chargeur', 'Powerbank']
        if not all(col in df.columns for col in required):
            messagebox.showerror(
                'Format invalide',
                "Le fichier doit contenir les colonnes 'N° Tablette', 'Chargeur' et 'Powerbank'.",
            )
            return

        if 'Statut actuel' not in df.columns:
            df['Statut actuel'] = 'En stock'
        if 'Observations' not in df.columns:
            df['Observations'] = 'RAS'

        # Remplit les observations manquantes
        df['Observations'] = df['Observations'].fillna('RAS').replace('', 'RAS')

        df = df[
            ['N° Tablette', 'Statut actuel', 'Chargeur', 'Powerbank', 'Observations']
        ]

        df.to_excel(TABLETTES_FILE, index=False, engine='openpyxl')
        messagebox.showinfo('Import', 'Fichier tablettes importé avec succès.')
        update_dashboard()


def affectation_automatique():
    """Réalise une affectation aléatoire depuis deux fichiers Excel."""
    ben_path = filedialog.askopenfilename(
        title='Fichier bénéficiaires',
        filetypes=[('Fichiers Excel', '*.xlsx')],
    )
    if not ben_path:
        return
    tab_path = filedialog.askopenfilename(
        title='Fichier tablettes',
        filetypes=[('Fichiers Excel', '*.xlsx')],
    )
    if not tab_path:
        return
    df_ben = pd.read_excel(ben_path, engine='openpyxl')
    df_tab = pd.read_excel(tab_path, engine='openpyxl')
    df_tab = df_tab[df_tab['Statut'] == 'En stock']
    if len(df_tab) < len(df_ben):
        messagebox.showerror('Erreur', "Nombre de tablettes insuffisant")
        return
    df_tab = df_tab.sample(frac=1).reset_index(drop=True).head(len(df_ben))
    df_ben = df_ben.sample(frac=1).reset_index(drop=True)
    df_aff = pd.concat(
        [
            df_ben.reset_index(drop=True),
            df_tab[['N° tablette', 'N° chargeur']].reset_index(drop=True),
        ],
        axis=1,
    )
    df_aff['Superviseur'] = ''
    df_aff['Téléphone superviseur'] = ''
    df_aff['Date affectation'] = datetime.today().strftime('%Y-%m-%d')
    df_aff = df_aff[
        [
            'Groupe',
            'Nom',
            'Fonction',
            'Téléphone',
            'N° tablette',
            'N° chargeur',
            'Superviseur',
            'Téléphone superviseur',
            'Date affectation',
        ]
    ]
    df_aff.rename(
        columns={'Nom': 'Bénéficiaire', 'N° chargeur': 'N° chargeur tablette'},
        inplace=True,
    )
    os.makedirs(FICHES_DIR, exist_ok=True)
    df_aff.to_excel(AFFECT_FILE, index=False, engine='openpyxl')
    for _, row in df_aff.iterrows():
        mapping = {
            'Groupe': row['Groupe'],
            'Bénéficiaire': row['Bénéficiaire'],
            'Fonction': row['Fonction'],
            'Téléphone': str(row['Téléphone']),
            'N° tablette': str(row['N° tablette']),
            'N° chargeur tablette': str(row['N° chargeur tablette']),
            'Superviseur': row['Superviseur'],
            'Téléphone superviseur': str(row['Téléphone superviseur']),
            'Date affectation': str(row['Date affectation']),
        }
        generer_fiche(mapping, row['Bénéficiaire'])
    messagebox.showinfo('Succès', 'Affectations générées.')
    update_dashboard()


def ajouter_tablette():
    num = entry_num_new.get().strip()
    if not num:
        messagebox.showwarning('Champs manquant', 'Veuillez entrer un numéro de tablette.')
        return
    df = load_tablettes()
    if num in df['N° Tablette'].astype(str).values:
        messagebox.showerror('Erreur', 'La tablette existe déjà.')
        return

    df.loc[len(df)] = {
        'N° Tablette': num,
        'Statut actuel': status_var.get(),
        'Chargeur': 'Oui' if charger_var.get() else 'Non',
        'Powerbank': 'Oui' if powerbank_var.get() else 'Non',
        'Observations': 'RAS',
    }
    save_tablettes(df)
    messagebox.showinfo('Succès', 'Tablette enregistrée.')
    update_dashboard()
    entry_num_new.delete(0, tk.END)
    charger_var.set(False)
    powerbank_var.set(False)
    status_var.set(STATUS_OPTIONS[0])


def assigner_tablette():
    num = entry_num_aff.get().strip()
    nom = entry_nom.get().strip()
    ident = entry_ident.get().strip()
    date = entry_date_aff.get().strip()
    if not num or not nom or not ident or not date:
        messagebox.showwarning('Champs manquants', 'Veuillez remplir tous les champs.')
        return
    df = load_tablettes()
    if num not in df['N° Tablette'].astype(str).values:
        messagebox.showerror('Erreur', "La tablette n'existe pas")
        return
    status = df.loc[df['N° Tablette'].astype(str) == num, 'Statut actuel'].iloc[0]
    if status != 'En stock':
        messagebox.showerror('Indisponible', 'La tablette n\'est pas en stock.')
        return
    df.loc[df['N° Tablette'].astype(str) == num, 'Statut actuel'] = 'Affectée'
    save_tablettes(df)

    df_aff = load_affectations()
    df_aff.loc[len(df_aff)] = {
        "Date d'affectation": date,
        'N° Tablette': num,
        'Nom bénéficiaire': nom,
        'Identifiant bénéficiaire': ident,
    }
    save_affectations(df_aff)
    generer_fiche(
        {
            'Groupe': '',
            'Bénéficiaire': nom,
            'Fonction': '',
            'Téléphone': '',
            'N° tablette': num,
            'N° chargeur tablette': '',
            'Superviseur': '',
            'Téléphone superviseur': '',
            'Date affectation': date,
        },
        nom,
    )
    messagebox.showinfo('Succès', 'Tablette affectée.')
    update_dashboard()
    entry_num_aff.delete(0, tk.END)
    entry_nom.delete(0, tk.END)
    entry_ident.delete(0, tk.END)
    entry_date_aff.delete(0, tk.END)
    entry_date_aff.insert(0, datetime.today().strftime('%Y-%m-%d'))


def retour_tablette():
    num = entry_num_retour.get().strip()
    if not num:
        messagebox.showwarning('Champs manquant', 'Veuillez entrer un numéro de tablette.')
        return
    df = load_tablettes()
    if num not in df['N° Tablette'].astype(str).values:
        messagebox.showerror('Erreur', "La tablette n'existe pas")
        return
    df.loc[df['N° Tablette'].astype(str) == num, 'Statut actuel'] = 'En stock'
    # Recherche le dernier incident lie a cette tablette pour mettre a jour
    df_inc = load_incidents()
    incidents_tablette = df_inc[df_inc['N° Tablette'].astype(str) == num]
    if not incidents_tablette.empty:
        last_nature = incidents_tablette.iloc[-1]['Nature incident']
        df.loc[df['N° Tablette'].astype(str) == num, 'Observations'] = last_nature
    save_tablettes(df)
    messagebox.showinfo('Succès', 'Tablette retournée en stock.')
    update_dashboard()
    entry_num_retour.delete(0, tk.END)


def declarer_incident():
    num = entry_num_inc.get().strip()
    nature = entry_nature.get().strip()
    declarant = entry_declarant.get().strip()
    lieu = entry_lieu.get().strip()
    date = entry_date_inc.get().strip()
    if not all([num, nature, declarant, lieu, date]):
        messagebox.showwarning('Champs manquants', 'Veuillez remplir tous les champs.')
        return
    df_inc = load_incidents()
    df_inc.loc[len(df_inc)] = {
        'Date': date,
        'N° Tablette': num,
        'Nature incident': nature,
        'Déclarant': declarant,
        'Lieu': lieu,
    }
    save_incidents(df_inc)

    # Met a jour l'observation de la tablette avec la nature de l'incident
    df_tab = load_tablettes()
    if num in df_tab['N° Tablette'].astype(str).values:
        df_tab.loc[df_tab['N° Tablette'].astype(str) == num, 'Observations'] = nature
        save_tablettes(df_tab)
    messagebox.showinfo('Succès', 'Incident déclaré.')
    update_dashboard()
    entry_num_inc.delete(0, tk.END)
    entry_nature.delete(0, tk.END)
    entry_declarant.delete(0, tk.END)
    entry_lieu.delete(0, tk.END)
    entry_date_inc.delete(0, tk.END)
    entry_date_inc.insert(0, datetime.today().strftime('%Y-%m-%d'))


def update_dashboard():
    df = load_tablettes()
    stock = (df['Statut actuel'] == 'En stock').sum()
    affect = (df['Statut actuel'] == 'Affectée').sum()
    df_inc = load_incidents()
    incidents = len(df_inc)
    label_stock.config(text=f'En stock : {stock}')
    label_affect.config(text=f'Affectées : {affect}')
    label_incident.config(text=f'Incidents : {incidents}')


# Initialisation des fichiers
init_files()

# Interface graphique
root = tk.Tk()
root.title('Gestion des tablettes')
root.geometry('650x520')

# Amélioration de l'apparence avec le thème ttk
style = ttk.Style(root)
style.theme_use('clam')
style.configure('.', font=('Helvetica', 12))
style.configure('TButton', font=('Helvetica', 12), padding=6)
style.configure('TEntry', font=('Helvetica', 12))
style.configure('TLabel', font=('Helvetica', 12))

notebook = ttk.Notebook(root)
notebook.pack(expand=1, fill='both')

# --- Onglet Affectation ---
aff_frame = ttk.Frame(notebook)
notebook.add(aff_frame, text='Affectation')

label_num_aff = ttk.Label(aff_frame, text='N° Tablette :')
label_num_aff.grid(row=0, column=0, sticky='e')
entry_num_aff = ttk.Entry(aff_frame, width=25)
entry_num_aff.grid(row=0, column=1, pady=2)

label_nom = ttk.Label(aff_frame, text='Nom bénéficiaire :')
label_nom.grid(row=1, column=0, sticky='e')
entry_nom = ttk.Entry(aff_frame, width=25)
entry_nom.grid(row=1, column=1, pady=2)

label_ident = ttk.Label(aff_frame, text='ID bénéficiaire :')
label_ident.grid(row=2, column=0, sticky='e')
entry_ident = ttk.Entry(aff_frame, width=25)
entry_ident.grid(row=2, column=1, pady=2)

label_date_aff = ttk.Label(aff_frame, text='Date :')
label_date_aff.grid(row=3, column=0, sticky='e')
entry_date_aff = ttk.Entry(aff_frame, width=25)
entry_date_aff.insert(0, datetime.today().strftime('%Y-%m-%d'))
entry_date_aff.grid(row=3, column=1, pady=2)

btn_affecter = ttk.Button(aff_frame, text='Affecter', command=assigner_tablette)
btn_affecter.grid(row=4, column=0, columnspan=2, pady=5)

btn_auto = ttk.Button(
    aff_frame,
    text='Affectation automatique',
    command=affectation_automatique,
)
btn_auto.grid(row=5, column=0, columnspan=2, pady=5)

btn_fiches = ttk.Button(
    aff_frame,
    text='Ouvrir dossier fiches',
    command=ouvrir_dossier_fiches,
)
btn_fiches.grid(row=6, column=0, columnspan=2, pady=5)

# --- Onglet Retour ---
retour_frame = ttk.Frame(notebook)
notebook.add(retour_frame, text='Retour')

label_num_retour = ttk.Label(retour_frame, text='N° Tablette :')
label_num_retour.grid(row=0, column=0, sticky='e')
entry_num_retour = ttk.Entry(retour_frame, width=25)
entry_num_retour.grid(row=0, column=1, pady=2)

btn_retour = ttk.Button(retour_frame, text='Enregistrer le retour', command=retour_tablette)
btn_retour.grid(row=1, column=0, columnspan=2, pady=5)

# --- Onglet Enregistrement ---
enreg_frame = ttk.Frame(notebook)
notebook.add(enreg_frame, text='Enregistrement')

label_num_new = ttk.Label(enreg_frame, text='N° Tablette :')
label_num_new.grid(row=0, column=0, sticky='e')
entry_num_new = ttk.Entry(enreg_frame, width=25)
entry_num_new.grid(row=0, column=1, pady=2)

charger_var = tk.BooleanVar()
check_charger = ttk.Checkbutton(enreg_frame, text='Chargeur présent', variable=charger_var)
check_charger.grid(row=1, column=0, columnspan=2, sticky='w', pady=2)

powerbank_var = tk.BooleanVar()
check_powerbank = ttk.Checkbutton(enreg_frame, text='Powerbank présente', variable=powerbank_var)
check_powerbank.grid(row=2, column=0, columnspan=2, sticky='w', pady=2)

# Sélection du statut initial de la tablette
status_var = tk.StringVar(value=STATUS_OPTIONS[0])
label_status = ttk.Label(enreg_frame, text='Statut actuel :')
label_status.grid(row=3, column=0, sticky='e')
combo_status = ttk.Combobox(
    enreg_frame,
    textvariable=status_var,
    values=STATUS_OPTIONS,
    state='readonly',
)
combo_status.grid(row=3, column=1, pady=2)
combo_status.current(0)

btn_add = ttk.Button(enreg_frame, text='Enregistrer', command=ajouter_tablette)
btn_add.grid(row=4, column=0, columnspan=2, pady=5)

btn_import = ttk.Button(
    enreg_frame,
    text='Importer depuis Excel',
    command=importer_tablettes,
)
btn_import.grid(row=5, column=0, columnspan=2, pady=5)

# --- Onglet Incident ---
incident_frame = ttk.Frame(notebook)
notebook.add(incident_frame, text='Incident')

label_num_inc = ttk.Label(incident_frame, text='N° Tablette :')
label_num_inc.grid(row=0, column=0, sticky='e')
entry_num_inc = ttk.Entry(incident_frame, width=25)
entry_num_inc.grid(row=0, column=1, pady=2)

label_nature = ttk.Label(incident_frame, text='Nature :')
label_nature.grid(row=1, column=0, sticky='e')
entry_nature = ttk.Entry(incident_frame, width=25)
entry_nature.grid(row=1, column=1, pady=2)

label_declarant = ttk.Label(incident_frame, text='Déclarant :')
label_declarant.grid(row=2, column=0, sticky='e')
entry_declarant = ttk.Entry(incident_frame, width=25)
entry_declarant.grid(row=2, column=1, pady=2)

label_lieu = ttk.Label(incident_frame, text='Lieu :')
label_lieu.grid(row=3, column=0, sticky='e')
entry_lieu = ttk.Entry(incident_frame, width=25)
entry_lieu.grid(row=3, column=1, pady=2)

label_date_inc = ttk.Label(incident_frame, text='Date :')
label_date_inc.grid(row=4, column=0, sticky='e')
entry_date_inc = ttk.Entry(incident_frame, width=25)
entry_date_inc.insert(0, datetime.today().strftime('%Y-%m-%d'))
entry_date_inc.grid(row=4, column=1, pady=2)

btn_incident = ttk.Button(incident_frame, text='Déclarer', command=declarer_incident)
btn_incident.grid(row=5, column=0, columnspan=2, pady=5)

# --- Onglet Bases de données ---
data_frame = ttk.Frame(notebook)
notebook.add(data_frame, text='Bases')

btn_open_tab = ttk.Button(data_frame, text='Tablettes', command=open_tablettes)
btn_open_tab.pack(pady=2, fill='x')
btn_open_aff = ttk.Button(data_frame, text='Affectations', command=open_affectations)
btn_open_aff.pack(pady=2, fill='x')
btn_open_inc = ttk.Button(data_frame, text='Incidents', command=open_incidents)
btn_open_inc.pack(pady=2, fill='x')

# --- Onglet Tableau de bord ---
db_frame = ttk.Frame(notebook)
notebook.add(db_frame, text='Tableau de bord')

label_stock = ttk.Label(db_frame, text='En stock : 0')
label_stock.pack(pady=2)
label_affect = ttk.Label(db_frame, text='Affectées : 0')
label_affect.pack(pady=2)
label_incident = ttk.Label(db_frame, text='Incidents : 0')
label_incident.pack(pady=2)

update_dashboard()
root.mainloop()
